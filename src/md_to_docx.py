import argparse
import re
import sys
from pathlib import Path
from urllib.parse import unquote

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DEFAULT_TEMPLATE_DIRS = [
    Path("data/example format"),
    Path("data/示例模板"),
    Path("data/示例模板文件夹"),
]

DEFAULT_AUTHOR = "王伟杰"
DEFAULT_CONTACT = "tjwanggroup@163.com"

IMAGE_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
DOI_URL_RE = re.compile(r"https?://(?:dx\.)?doi\.org/[^\s\]\)\"'，。；;]+", re.I)
DOI_RE = re.compile(r"\b10\.\d{4,9}/[-._;()/:A-Z0-9]+", re.I)


def find_template(template_arg=None):
    if template_arg:
        template = Path(template_arg)
        if not template.exists():
            raise FileNotFoundError(f"Template not found: {template}")
        return template

    for template_dir in DEFAULT_TEMPLATE_DIRS:
        if template_dir.exists():
            templates = sorted(template_dir.glob("*.docx"))
            if templates:
                return templates[0]
    raise FileNotFoundError(
        "No .docx template found under data/example format or data/示例模板*."
    )


def set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=12, bold=None):
    run.font.name = ascii_font
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)


def set_style_font(style, east_asia="宋体", ascii_font="Times New Roman", size=12, bold=None):
    style.font.name = ascii_font
    style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold

    r_pr = style.element.rPr
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        style.element.append(r_pr)
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)


def configure_template_like_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)

    normal = doc.styles["Normal"]
    set_style_font(normal, east_asia="宋体", ascii_font="Times New Roman", size=12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(0.8467)
    normal.paragraph_format.line_spacing = Pt(20)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    heading1 = doc.styles["Heading 1"]
    set_style_font(heading1, east_asia="黑体", ascii_font="Times New Roman", size=14, bold=True)
    heading1.paragraph_format.first_line_indent = Cm(0)
    heading1.paragraph_format.space_before = Pt(0)
    heading1.paragraph_format.space_after = Pt(0)

    heading2 = doc.styles["Heading 2"]
    set_style_font(heading2, east_asia="黑体", ascii_font="Times New Roman", size=12, bold=True)
    heading2.paragraph_format.first_line_indent = Cm(0)
    heading2.paragraph_format.space_before = Pt(0)
    heading2.paragraph_format.space_after = Pt(0)

    title = doc.styles["Title"]
    set_style_font(title, east_asia="宋体", ascii_font="Times New Roman", size=10.5)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Cm(0)
    title.paragraph_format.line_spacing = Pt(12)
    title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST


def clear_document_body(doc):
    body = doc._body._element
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def clean_math_expression(expr):
    replacements = {
        r"\cdot": "·",
        r"\times": "×",
        r"\geq": "≥",
        r"\leq": "≤",
        r"\Delta": "Δ",
        r"\alpha": "α",
        r"\beta": "β",
        r"\gamma": "γ",
        r"\lambda": "λ",
    }
    for old, new in replacements.items():
        expr = expr.replace(old, new)
    expr = re.sub(r"\\mathrm\s*\{([^{}]*)\}", r"\1", expr)
    expr = expr.replace(r"\left", "").replace(r"\right", "")
    expr = re.sub(r"\\([A-Za-z]+)", r"\1", expr)
    expr = expr.replace("{", "").replace("}", "")
    expr = re.sub(r"\s+", " ", expr)
    return expr.strip()


def clean_inline_text(text):
    text = re.sub(r"\$([^$]+)\$", lambda m: clean_math_expression(m.group(1)), text)
    text = text.replace("**", "")
    text = text.replace("__", "")
    text = text.replace("`", "")
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = text.replace(r"\_", "_")
    return text.strip()


RICH_TOKEN_RE = re.compile(
    r"Δδ18O|δ18O|NH4\+|NO[23]-|N2O|H2O|CO2|CH4|O2|m3"
)


def rich_text_segments(text):
    position = 0
    for match in RICH_TOKEN_RE.finditer(text):
        if match.start() > position:
            yield text[position : match.start()], None

        token = match.group(0)
        if token in {"Δδ18O", "δ18O"}:
            prefix = token[:-3]
            yield prefix, None
            yield "18", "superscript"
            yield "O", None
        elif token == "NH4+":
            yield "NH", None
            yield "4", "subscript"
            yield "+", "superscript"
        elif token.startswith("NO") and token.endswith("-"):
            yield "NO", None
            yield token[2], "subscript"
            yield "-", "superscript"
        elif token == "m3":
            yield "m", None
            yield "3", "superscript"
        else:
            # Binary formulas such as N2O, O2, H2O, CO2 and CH4.
            for char in token:
                yield char, "subscript" if char.isdigit() else None

        position = match.end()

    if position < len(text):
        yield text[position:], None


def normalize_caption(caption):
    caption = clean_inline_text(caption)
    caption = re.sub(r"^(图\s*\d+)\s*[：:]\s*", r"\1 ", caption)
    return caption


def parse_image_target(raw_target):
    target = raw_target.strip()
    if target.startswith("<") and ">" in target:
        target = target[1 : target.index(">")]
    else:
        title_match = re.match(r"(.+?)(?:\s+[\"'][^\"']+[\"'])?$", target)
        if title_match:
            target = title_match.group(1)
    return unquote(target.strip())


def extract_titles(markdown):
    h1_titles = []
    for line in markdown.splitlines():
        match = re.match(r"^#\s+(.+?)\s*$", line)
        if match:
            h1_titles.append(clean_inline_text(match.group(1)))
        if len(h1_titles) >= 2:
            break
    title_en = h1_titles[0] if h1_titles else "Untitled"
    title_zh = h1_titles[1] if len(h1_titles) > 1 else ""
    return title_en, title_zh


def parse_markdown_blocks(markdown):
    blocks = []
    paragraph_lines = []
    skipped_top_h1 = 0

    def flush_paragraph():
        if paragraph_lines:
            text = clean_inline_text(" ".join(line.strip() for line in paragraph_lines))
            if text and not text.startswith("原文信息："):
                blocks.append({"type": "paragraph", "text": text})
            paragraph_lines.clear()

    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            flush_paragraph()
            continue

        heading = HEADING_RE.match(line)
        if heading:
            flush_paragraph()
            level = len(heading.group(1))
            text = clean_inline_text(heading.group(2))
            if level == 1 and skipped_top_h1 < 2:
                skipped_top_h1 += 1
                continue
            if text.startswith("主要内容"):
                text = "主要内容"
            if level >= 3:
                text = re.sub(r"^(\d+)\.\s+", r"\1 ", text)
            blocks.append({"type": "heading", "level": level, "text": text})
            continue

        image = IMAGE_RE.match(line)
        if image:
            flush_paragraph()
            blocks.append(
                {
                    "type": "image",
                    "caption": normalize_caption(image.group(1)),
                    "target": parse_image_target(image.group(2)),
                }
            )
            continue

        paragraph_lines.append(line)

    flush_paragraph()
    return blocks


def markdown_image_targets(markdown):
    targets = set()
    for raw_target in re.findall(r"!\[[^\]]*\]\(([^)]+)\)", markdown):
        target = parse_image_target(raw_target)
        if target.startswith("images/"):
            targets.add(target)
    return targets


def strip_trailing_doi_punctuation(value):
    return value.rstrip(".,;:，。；：)")


def extract_doi_text(paths):
    for path in paths:
        if not path or not path.exists():
            continue
        text = path.read_text(encoding="utf-8", errors="ignore")
        url_match = DOI_URL_RE.search(text)
        if url_match:
            return strip_trailing_doi_punctuation(url_match.group(0))
        doi_match = DOI_RE.search(text)
        if doi_match:
            return "https://doi.org/" + strip_trailing_doi_punctuation(doi_match.group(0))
    return ""


def extract_authors(summary_text, full_md_path):
    if full_md_path.exists():
        lines = [line.strip() for line in full_md_path.read_text(encoding="utf-8", errors="ignore").splitlines()]
        non_empty = [line for line in lines if line]
        if len(non_empty) >= 2:
            candidate = clean_inline_text(non_empty[1])
            candidate = candidate.replace("\\$", "").replace("\\*", "")
            candidate = re.sub(r"\s+", " ", candidate)
            candidate = re.sub(r",(?=\S)", ", ", candidate)
            if len(candidate) <= 240:
                return candidate

    match = re.search(r"^原文信息：(.+)$", summary_text, re.M)
    if match:
        candidate = clean_inline_text(match.group(1))
        candidate = re.split(r"\.\s+\*|,?\s+\*[A-Z]", candidate, maxsplit=1)[0]
        return candidate.strip()
    return ""


def image_is_reasonable_graphical_abstract(path):
    try:
        from PIL import Image

        with Image.open(path) as image:
            width, height = image.size
    except Exception:
        return False

    if width < 300 or height < 150:
        return False
    ratio = width / max(height, 1)
    return 0.8 <= ratio <= 5.5


def looks_like_graphical_abstract_text(text):
    text = text.lower()
    keywords = [
        "graphical abstract",
        "toc graphic",
        "table of contents graphic",
        "图文摘要",
        "图形摘要",
    ]
    return any(keyword in text for keyword in keywords)


def find_graphical_abstract(input_dir, image_dir, summary_text):
    used_targets = markdown_image_targets(summary_text)

    full_md = input_dir / "full.md"
    if full_md.exists():
        lines = full_md.read_text(encoding="utf-8", errors="ignore").splitlines()
        for index, line in enumerate(lines):
            image = IMAGE_RE.match(line.strip())
            if not image:
                continue
            target = parse_image_target(image.group(2))
            nearby = "\n".join(lines[max(0, index - 3) : min(len(lines), index + 4)])
            if target in used_targets or not looks_like_graphical_abstract_text(nearby):
                continue
            image_path = resolve_image_path(target, image_dir, full_md)
            if image_path.exists() and image_is_reasonable_graphical_abstract(image_path):
                return {"target": target, "path": image_path, "caption": "图文摘要"}

    for json_path in sorted(input_dir.glob("*.json")):
        try:
            import json

            data = json.loads(json_path.read_text(encoding="utf-8", errors="ignore"))
        except Exception:
            continue
        if not isinstance(data, list):
            continue

        before_main_text = True
        for item in data:
            text = str(item.get("text", "") or item.get("caption", "") or "")
            normalized = text.upper()
            if "INTRODUCTION" in normalized or "引言" in text:
                before_main_text = False

            if item.get("type") != "image":
                continue
            target = item.get("img_path") or item.get("image_path") or item.get("path")
            if not target or target in used_targets:
                continue
            image_path = image_dir / target.split("images/", 1)[-1]
            caption_text = str(item.get("caption", "") or text)
            is_named_graphical_abstract = looks_like_graphical_abstract_text(caption_text)
            if not (before_main_text or is_named_graphical_abstract):
                continue
            if image_path.exists() and image_is_reasonable_graphical_abstract(image_path):
                return {"target": target, "path": image_path, "caption": "图文摘要"}

    return None


def add_graphical_abstract(doc, graphical_abstract):
    add_picture(doc, graphical_abstract["path"])
    add_caption(doc, graphical_abstract.get("caption") or "图文摘要")


def find_input_cover_image(image_dir):
    if not image_dir.exists():
        return None

    exact_names = [
        "封面.png",
        "封面.jpg",
        "封面.jpeg",
        "cover.png",
        "cover.jpg",
        "cover.jpeg",
        "Cover.png",
        "Cover.jpg",
        "Cover.jpeg",
    ]
    for name in exact_names:
        path = image_dir / name
        if path.exists():
            return path

    for path in sorted(image_dir.iterdir()):
        if path.is_file() and path.suffix.lower() in {".png", ".jpg", ".jpeg"}:
            if "封面" in path.stem or "cover" in path.stem.lower():
                return path
    return None


def add_heading(doc, text, level):
    style = "Heading 1" if level <= 2 else "Heading 2"
    if level >= 3:
        style = "Heading 2"
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    add_rich_text(
        paragraph,
        text,
        east_asia="黑体",
        ascii_font="Times New Roman",
        size=14 if style == "Heading 1" else 12,
        bold=True,
    )
    return paragraph


def add_rich_text(paragraph, text, east_asia="宋体", ascii_font="Times New Roman", size=12, bold=None):
    for segment, vertical_align in rich_text_segments(text):
        if not segment:
            continue
        run = paragraph.add_run(segment)
        set_run_font(run, east_asia=east_asia, ascii_font=ascii_font, size=size, bold=bold)
        if vertical_align == "subscript":
            run.font.subscript = True
        elif vertical_align == "superscript":
            run.font.superscript = True
    return paragraph


def add_normal_paragraph(doc, text, first_line=True):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(0.8467) if first_line else Cm(0)
    paragraph.paragraph_format.line_spacing = Pt(20)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    add_rich_text(paragraph, text, east_asia="宋体", ascii_font="Times New Roman", size=12)
    return paragraph


def add_caption(doc, caption):
    paragraph = doc.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = Pt(12)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    add_rich_text(paragraph, caption, east_asia="宋体", ascii_font="Times New Roman", size=10.5)
    return paragraph


def add_picture(doc, image_path, width_cm=14.65):
    paragraph = doc.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    return paragraph


def resolve_image_path(target, image_dir, md_file):
    if target.startswith("images/"):
        return image_dir / target.split("images/", 1)[1]
    path = Path(target)
    if path.is_absolute():
        return path
    return md_file.parent / path


def get_font(paths, size):
    from PIL import ImageFont

    for path in paths:
        if Path(path).exists():
            return ImageFont.truetype(path, size=size)
    return ImageFont.load_default(size=size)


def draw_wrapped_text(draw, xy, text, font, fill, max_width, line_gap=8):
    lines = wrap_text_lines(draw, text, font, max_width)
    return draw_lines(draw, xy, lines, font, fill, line_gap)


def wrap_text_lines(draw, text, font, max_width):
    words = text.split() or [text]
    lines = []
    current = ""

    def append_token(token):
        token_lines = []
        current_token = ""
        for char in token:
            trial = f"{current_token}{char}"
            if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
                current_token = trial
            else:
                if current_token:
                    token_lines.append(current_token)
                current_token = char
        if current_token:
            token_lines.append(current_token)
        return token_lines

    for word in words:
        if draw.textbbox((0, 0), word, font=font)[2] > max_width:
            if current:
                lines.append(current)
                current = ""
            lines.extend(append_token(word))
            continue
        trial = f"{current} {word}".strip()
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_lines(draw, xy, lines, font, fill, line_gap=8):
    x, y = xy
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        bbox = draw.textbbox((x, y), line, font=font)
        y = bbox[3] + line_gap
    return y


def get_font_from_candidates(paths, size):
    from PIL import ImageFont

    for path in paths:
        if Path(path).exists():
            return ImageFont.truetype(path, size=size)
    return ImageFont.load_default(size=size)


def fit_font_and_lines(draw, text, font_paths, max_width, max_height, start_size, min_size, line_gap=8):
    for size in range(start_size, min_size - 1, -4):
        font = get_font_from_candidates(font_paths, size)
        lines = wrap_text_lines(draw, text, font, max_width)
        if not lines:
            return font, []
        heights = []
        for line in lines:
            bbox = draw.textbbox((0, 0), line, font=font)
            heights.append(bbox[3] - bbox[1])
        total_height = sum(heights) + max(0, len(lines) - 1) * line_gap
        if total_height <= max_height:
            return font, lines
    font = get_font_from_candidates(font_paths, min_size)
    return font, wrap_text_lines(draw, text, font, max_width)


def create_cover_image(paper_id, title_en, title_zh, doi_url, authors, output_dir):
    try:
        from PIL import Image, ImageDraw
    except Exception:
        return None

    output_dir.mkdir(parents=True, exist_ok=True)
    cover_path = output_dir / f"{paper_id}_cover.png"
    image = Image.new("RGB", (2480, 1080), "white")
    draw = ImageDraw.Draw(image)

    regular_paths = [r"C:\Windows\Fonts\arial.ttf", r"C:\Windows\Fonts\msyh.ttc"]
    regular_cn_paths = [r"C:\Windows\Fonts\msyh.ttc", r"C:\Windows\Fonts\arial.ttf"]
    bold_paths = [r"C:\Windows\Fonts\arialbd.ttf", r"C:\Windows\Fonts\msyhbd.ttc"]

    font_regular = get_font(
        [r"C:\Windows\Fonts\arial.ttf", r"C:\Windows\Fonts\msyh.ttc"], 54
    )
    font_small = get_font(
        [r"C:\Windows\Fonts\arial.ttf", r"C:\Windows\Fonts\msyh.ttc"], 42
    )
    font_label = get_font(
        [r"C:\Windows\Fonts\arialbd.ttf", r"C:\Windows\Fonts\msyhbd.ttc"], 58
    )

    draw.text((110, 70), "literature | note", font=font_regular, fill=(30, 30, 30))
    draw.text((2120, 78), "Article", font=font_label, fill=(0, 0, 0))
    draw.rectangle((110, 170, 2370, 185), fill=(239, 67, 53))
    draw.text((110, 215), "Generated from MinerU-extracted academic paper", font=font_small, fill=(0, 0, 0))

    if doi_url:
        draw.text((1570, 410), doi_url, font=font_small, fill=(0, 102, 170))
    draw.line((110, 475, 2370, 475), fill=(0, 0, 0), width=4)

    font_bold, title_lines = fit_font_and_lines(
        draw, title_en, bold_paths, max_width=2140, max_height=300, start_size=104, min_size=64, line_gap=12
    )
    y = draw_lines(draw, (110, 520), title_lines, font_bold, (0, 0, 0), 12)
    if title_zh:
        y += 14
        remaining_height = max(64, 900 - y)
        font_regular_cn, zh_lines = fit_font_and_lines(
            draw, title_zh, regular_cn_paths, max_width=2140, max_height=remaining_height, start_size=48, min_size=34, line_gap=8
        )
        y = draw_lines(draw, (110, y), zh_lines, font_regular_cn, (0, 0, 0), 8)
    separator_y = min(max(y + 16, 930), 970)
    draw.line((110, separator_y, 2370, separator_y), fill=(0, 0, 0), width=4)
    if authors:
        author_y = separator_y + 18
        font_author, author_lines = fit_font_and_lines(
            draw, authors, regular_paths, max_width=2140, max_height=max(48, 1060 - author_y), start_size=34, min_size=24, line_gap=4
        )
        draw_lines(draw, (110, author_y), author_lines, font_author, (0, 0, 0), 4)

    image.save(cover_path)
    return cover_path


def convert_one(paper_id, args):
    md_file = Path(args.markdown) if args.markdown else Path("data/output") / f"{paper_id}_summary.md"
    if not md_file.exists():
        raise FileNotFoundError(f"Markdown summary not found: {md_file}")

    output_file = Path(args.output) if args.output else Path("data/output") / f"{paper_id}_summary.docx"
    input_dir = Path("data/input") / paper_id
    image_dir = Path(args.images) if args.images else input_dir / "images"
    # The template is used as a formatting reference only. Creating a fresh
    # document avoids carrying hidden media such as the template paper cover.
    find_template(args.template)

    summary_text = md_file.read_text(encoding="utf-8")
    title_en, title_zh = extract_titles(summary_text)
    doi_url = args.doi or extract_doi_text(
        [md_file, input_dir / "full.md", *sorted(input_dir.glob("*.json"))]
    )
    authors = args.paper_authors or extract_authors(summary_text, input_dir / "full.md")

    doc = Document()
    configure_template_like_document(doc)
    doc.core_properties.title = title_zh or title_en
    graphical_abstract = None if args.no_graphical_abstract else find_graphical_abstract(input_dir, image_dir, summary_text)
    input_cover_path = None if args.no_cover else find_input_cover_image(image_dir)

    if input_cover_path:
        add_picture(doc, input_cover_path)
    elif args.generate_cover and not args.no_cover:
        cover_path = create_cover_image(
            paper_id, title_en, title_zh, doi_url, authors, Path(".tmp/docx_cover")
        )
        if cover_path:
            add_picture(doc, cover_path)

    add_heading(doc, "题目", 1)
    add_normal_paragraph(doc, title_en, first_line=False)
    if title_zh:
        add_normal_paragraph(doc, title_zh, first_line=False)

    in_abstract = False
    graphical_abstract_inserted = False
    for block in parse_markdown_blocks(summary_text):
        if (
            block["type"] == "heading"
            and in_abstract
            and block["text"] != "摘要"
            and graphical_abstract
            and not graphical_abstract_inserted
        ):
            add_graphical_abstract(doc, graphical_abstract)
            graphical_abstract_inserted = True
            in_abstract = False

        if block["type"] == "heading":
            add_heading(doc, block["text"], block["level"])
            in_abstract = block["text"] == "摘要"
        elif block["type"] == "paragraph":
            add_normal_paragraph(doc, block["text"], first_line=True)
        elif block["type"] == "image":
            image_path = resolve_image_path(block["target"], image_dir, md_file)
            if not image_path.exists():
                raise FileNotFoundError(f"Image referenced by markdown not found: {image_path}")
            add_picture(doc, image_path)
            if block["caption"]:
                add_caption(doc, block["caption"])

    if in_abstract and graphical_abstract and not graphical_abstract_inserted:
        add_graphical_abstract(doc, graphical_abstract)

    add_normal_paragraph(doc, "说明：本推送仅为学术交流，如有侵权，请联系删除。", first_line=False)
    if doi_url:
        add_normal_paragraph(doc, f"原文链接：{doi_url}", first_line=False)
    add_normal_paragraph(doc, f"投稿&合作请联系：{args.contact}", first_line=False)
    add_normal_paragraph(doc, f"作者：{args.author}", first_line=False)

    output_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_file)
    print(f"OK: wrote {output_file}")
    return output_file


def parse_args(argv):
    parser = argparse.ArgumentParser(
        description="Convert generated literature markdown summaries to template-styled DOCX."
    )
    parser.add_argument("paper_id", nargs="+", help="Paper id under data/input and data/output.")
    parser.add_argument("--template", help="Path to the reference .docx template.")
    parser.add_argument("--markdown", help="Override markdown input path; only for one paper id.")
    parser.add_argument("--images", help="Override images directory; only for one paper id.")
    parser.add_argument("--output", help="Override docx output path; only for one paper id.")
    parser.add_argument("--doi", help="Override DOI URL written to the output.")
    parser.add_argument("--paper-authors", help="Override paper authors written to the generated cover image.")
    parser.add_argument("--author", default=DEFAULT_AUTHOR, help="Reading note author footer.")
    parser.add_argument("--contact", default=DEFAULT_CONTACT, help="Contact footer.")
    parser.add_argument("--generate-cover", action="store_true", help="Generate an article-info cover image. Disabled by default.")
    parser.add_argument("--no-cover", action="store_true", help=argparse.SUPPRESS)
    parser.add_argument("--no-graphical-abstract", action="store_true", help="Do not auto-insert a graphical abstract image after the abstract.")
    args = parser.parse_args(argv)

    if len(args.paper_id) > 1 and (args.markdown or args.images or args.output):
        parser.error("--markdown, --images, and --output can only be used with one paper id.")
    return args


def main(argv=None):
    args = parse_args(argv or sys.argv[1:])
    status = 0
    for paper_id in args.paper_id:
        try:
            convert_one(paper_id, args)
        except Exception as exc:
            print(f"ERROR: {paper_id}: {exc}", file=sys.stderr)
            status = 1
    return status


if __name__ == "__main__":
    sys.exit(main())
